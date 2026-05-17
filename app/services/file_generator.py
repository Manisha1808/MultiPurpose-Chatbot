from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import getSampleStyleSheet

from docx import Document
import os


def generate_pdf(content, filename):

    os.makedirs("generated_files", exist_ok=True)

    file_path = (
        f"generated_files/{filename}.pdf"
    )

    doc = SimpleDocTemplate(file_path)

    styles = getSampleStyleSheet()

    story = []

    paragraphs = content.split("\n")

    for para in paragraphs:

        story.append(
            Paragraph(para, styles["BodyText"])
        )

        story.append(Spacer(1, 12))

    doc.build(story)

    return file_path


def generate_docx(content, filename):

    os.makedirs("generated_files", exist_ok=True)

    file_path = (
        f"generated_files/{filename}.docx"
    )

    doc = Document()

    doc.add_heading("Generated Document", level=1)

    doc.add_paragraph(content)

    doc.save(file_path)

    return file_path